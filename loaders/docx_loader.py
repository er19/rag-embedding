from docx import Document as DocxDocument

from .base import BaseLoader, Document


class DocxLoader(BaseLoader):
    def load(self, file_path: str) -> list[Document]:
        doc = DocxDocument(file_path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        return [Document(content=text, metadata={"source": file_path})]
